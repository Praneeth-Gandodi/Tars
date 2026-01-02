import os
import json
import subprocess
import shutil
import platform
from pathlib import Path
from docx import Document
from pypdf import PdfReader
from rich.console import Console 

console = Console()

homedir = Path.home()

def list_files_in_directory(directory):
    """
    List of files in the directory
    Args:
        directory (str): Returns the list of files in that directory

    Returns:
        list (list): returns the list of files in the directory.
    """
    dir = homedir / directory
    dir_path = Path(dir)
    if not dir_path.exists():
        return f"File not find in the directory you specified {directory}"
    return [item.name for item in dir_path.iterdir()]
 
def list_files_by_types(directory:str , ext:list):
    """
    List the files in that directory with the file type passed
    Args:
        directory (str): _description_
        ext (str): _description_
    Returns:
        list (list): Returns a list which contains the name of files with that type.
    """
    directory = homedir / directory
    dir_path = Path(directory)
    if not dir_path.exists():
        return "The file path doesnot Exits"
    files = []
    for extension in ext:
        extension = extension.lower().strip()
        extension = f"*.{extension}"
        files.extend([item.name for item in dir_path.glob(extension)])
    return files

def read_file_content(directory:str):
    """
    Returns the file content of the file.
    Args:
        directory (str): location of the file you want to open with the file name and extension.
    Returns:
        text (str): returns the content in the file. If file not found returns FileNotExist. 
    """
    dir_path = homedir / directory
    dir = Path(dir_path)
    supported_files = [".txt",".md",".csv",".log",".ini",".cfg",".env",".py",".html",".css",".js",".xml",".yaml",".yml"]

    if not dir.exists():
        return f"File path does not exists {dir}"
    
    file_type = Path(dir_path).suffix
    if file_type.strip() in supported_files:
        content = dir.read_text()
    elif file_type.strip() == ".json":
        try:
            with dir.open("r") as f:
                content = json.load(f)
        except Exception as e:
            return f"Error Occured while reading the file: {e}"
    elif file_type in ".docx":
        try:
            dir = Document(dir)
            return "\n".join(p.text for p in dir.paragraphs if p.text.strip())
        except Exception as e:
            return f"Error: Unable to read the file {e}"
    elif file_type == ".pdf":
        try:
            reader = PdfReader(dir)
            return "\n".join(page.extract_text() or "" for page in reader.pages)
        except Exception as e:
            return f"Error: Unable to read the PDF file: {e}"
    else:
        return f"File type not supported"
    return content

from pathlib import Path

def write_to_files(directory: str, content: str, mode: str = "w"):
    """
    Write or append data to a file.
    
    Args:
        directory (str): Full path to the file to write content.
        content (str): Text content to write or append.
        mode (str): 'w' for overwrite, 'a' for append. Default is 'w'.
        
    Returns:
        str: Status message describing the result.
            - "Created the file {filename} in {directory} successfully"
            - "Appended to file {filename} in {directory} successfully"
            - "Written to file {filename} in {directory} successfully"
            - "Error: {exception}" on failure
    """
    try:
        directory = homedir / directory
        file_path = Path(directory)
        file_path.parent.mkdir(parents=True, exist_ok=True)  
        if mode == "a" and file_path.exists():
            with file_path.open("a", encoding="utf-8") as f:
                f.write(content)
            return f"Appended to file {file_path.name} in {file_path.parent} successfully"
        
        else:  
            created_new = not file_path.exists()
            with file_path.open(mode, encoding="utf-8") as f:
                f.write(content)
            
            if created_new:
                return f"Created the file {file_path.name} in {file_path.parent} successfully"
            else:
                return f"Written to file {file_path.name} in {file_path.parent} successfully"
    
    except Exception as e:
        return f"Error: {e}"

from pathlib import Path
from docx import Document

def write_docx(directory: str, content: str, append:str):
    """
    Write or append content to a DOCX file.
    """
    directory = homedir / directory
    file_path = Path(directory)
    try:
        path = Path(file_path)
        path.parent.mkdir(parents=True, exist_ok=True)

        if path.exists() and append == "w":
            doc = Document(path)
        else:
            doc = Document() 

        doc.add_paragraph(content)
        doc.save(path)

        if not path.exists():
            return f"Created the DOCX file {path.name} in {path.parent} successfully"
        elif append:
            return f"Appended to DOCX file {path.name} in {path.parent} successfully"
        else:
            return f"Written to DOCX file {path.name} in {path.parent} successfully"

    except Exception as e:
        return f"Error: {e}"
    
def recursive_file_search(filename: str, start_dir: str = None):
    """
    Recursively search for a file by name or partial name.

    Args:
        filename (str): Name or partial name of the file to search.
        start_dir (str): Directory to start the search. Defaults to home directory.

    Returns:
        list: List of matching file paths. Empty list if nothing found.
    """
    start_dir = homedir / start_dir
    start_path = Path(start_dir) if start_dir else Path.home()
    if not start_path.exists():
        return f"Start directory does not exist: {start_path}"

    filename_lower = filename.lower()
    matches = []

    for path in start_path.rglob("*"):
        if path.is_file() and filename_lower in path.name.lower():
            matches.append(str(path))

    if matches:
        return f"Found these files in the given directory {matches}"
    else:
        return f"No files found matching '{filename}' in {start_path}"   


def open_file(path):
    path = homedir / path
    path = Path(path)
    if not path.exists():
        return f"File does not exist: {path}"

    system = platform.system()
    try:
        if system == "Windows":
            os.startfile(path)  # opens file with default app
        elif system == "Darwin":  # macOS
            subprocess.run(["open", str(path)])
        else:  # Linux
            subprocess.run(["xdg-open", str(path)])
        return f"Opened {path}"
    except Exception as e:
        return f"Error opening {path}: {e}"
    